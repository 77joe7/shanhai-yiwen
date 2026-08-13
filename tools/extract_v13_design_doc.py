"""Extract the supplied V1.3 design document into an auditable UTF-8 text file."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "山海异闻录_天地未定_游戏设计说明书暨开发任务书_V1.3_优化修订版.docx"
OUTPUT = ROOT / "剧情" / "第一卷_黑雨" / "第一章_黑雨" / "资料" / "V1.3说明书_提取摘录.txt"


def main() -> None:
    document = Document(SOURCE)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    OUTPUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"Extracted {len(lines)} blocks to {OUTPUT}")


if __name__ == "__main__":
    main()
